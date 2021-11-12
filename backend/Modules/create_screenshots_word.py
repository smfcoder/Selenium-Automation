import docx
from docx.shared import Inches

class wordDocument:
    def __init__(self):
        self.doc = None
    def openWordDocument(self,document_heading):
        self.doc = docx.Document()
        self.doc.add_heading(document_heading, 0)
    def insertImage(self,image_heading,path_of_image):
    
        try:
            self.doc.add_heading(image_heading, 3)
            self.doc.add_picture(path_of_image, width=Inches(6), height=Inches(4))
            return(True)
        except:
            print("Could not insert image")
            return(False)
            
    def saveWordDocument(self,name_of_document):
        self.doc.save('./Reports/'+name_of_document+'.docx')
